#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère la version corrigée du Livrable 1 (Référentiel de données — CHU Big Data).

Corrections appliquées (vérifiées contre les sources réelles + le DDL implémenté) :
  - §2.1  : les 8 besoins reproduits VERBATIM depuis l'énoncé (fin des besoins inventés)
  - §4    : volumétrie décès corrigée 25M (au lieu de 50M) ; hospi = 2 479 séjours
  - §5    : 6 dimensions alignées sur le DDL réel (dont Dim_Geographie) ; patient_id STRING ; sexe M/F
  - §6    : 4 faits alignés sur sql/ddl/02_faits.hql (clés conformes, mesures réelles,
            décès dégénéré région/sexe/âge, satisfaction avec geo_id, hospi durée directe)
  - §8    : jobs ETL réalistes (pas de matching hospi<->décès, durée séjour donnée, pas de type_séjour)
  - §9    : sel par patient (linkage inter-années cohérent) ; pseudo réversible via mapping
  - §10   : NFR présentés comme cibles de conception, pas comme acquis
  - §12   : glossaire sans PySpark (pipeline 100% HiveQL)

Périmètre strict L1 : (A) modélisation axes + mesures, (B) jobs d'alimentation, (C) architecture.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/juliankerignard/Documents/CESI/BlocBigData/Projet/Livrable 1 Groupe - v2 corrigé.docx"

# Palette
NAVY = RGBColor(0x1F, 0x35, 0x55)
BLUE = RGBColor(0x2E, 0x5E, 0xAA)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3555"
ALT_FILL = "EEF2F8"

doc = Document()

# --- Styles de base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 16, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, BLUE)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
        set_cell_bg(hdr[i], HEADER_FILL)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if ridx % 2 == 1:
                set_cell_bg(cells[i], ALT_FILL)
    if widths:
        for col, w in zip(t.columns, widths):
            for c in col.cells:
                c.width = Inches(w)
    doc.add_paragraph()
    return t


def bullets(items, style="List Bullet"):
    for it in items:
        doc.add_paragraph(it, style=style)


def spacer():
    doc.add_paragraph()


# =========================================================================
# PAGE DE GARDE
# =========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PROJET BIG DATA — ENTREPÔT DE DONNÉES SANTÉ CHU")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Livrable 1 : Référentiel de Données")
r.font.size = Pt(15)
r.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Modèle conceptuel  ·  Architecture médaillon  ·  Jobs d'alimentation ETL")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

spacer(); spacer()

add_table(
    ["Rôle", "Membre", "Fait principal", "Source"],
    [
        ["Contributeur", "Julian Kerignard", "Fait_Consultation", "PostgreSQL"],
        ["Contributeur", "Chloé Lagocki", "Fait_Hospitalisation", "PostgreSQL / CSV"],
        ["Contributeur", "Matthieu Michel", "Fait_Satisfaction", "Fichiers plats (eSATIS)"],
        ["Contributeur", "Maxime Auchet", "Fait_Deces", "Fichiers INSEE"],
    ],
    widths=[1.3, 2.0, 2.0, 1.7],
)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Projet : CHU Big Data    ·    Version 2.0 (corrigée)    ·    05 juin 2026")
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_page_break()

# =========================================================================
# 1. RÉSUMÉ EXÉCUTIF
# =========================================================================
doc.add_heading("1. Résumé exécutif", level=1)
doc.add_paragraph(
    "Ce rapport constitue le Livrable 1 du projet CHU Big Data : le référentiel de données de "
    "l'entrepôt décisionnel. Conformément à l'énoncé, il couvre strictement trois objets :"
)
bullets([
    "la modélisation des axes d'analyse et des mesures (modèle dimensionnel en constellation) ;",
    "le développement des jobs d'alimentation du schéma décisionnel (chaîne ETL) ;",
    "la description de l'architecture de l'entrepôt de données (médaillon Bronze/Silver/Gold).",
])
doc.add_paragraph(
    "Le modèle répond aux 8 besoins d'analyse imposés par le CHU (section 2.1). Il s'appuie sur "
    "4 tables de faits (Consultation, Hospitalisation, Satisfaction, Décès) et 6 dimensions conformes "
    "partagées (Temps, Patient, Diagnostic, Établissement, Professionnel, Géographie). "
    "L'implémentation physique, les tests et l'optimisation relèvent du Livrable 2 ; les tableaux de "
    "bord et le storytelling, du Livrable 3."
)

# =========================================================================
# 2. INTRODUCTION ET CONTEXTE
# =========================================================================
doc.add_heading("2. Introduction et contexte", level=1)

doc.add_heading("2.1 Objectif du projet — les 8 besoins d'analyse", level=2)
doc.add_paragraph(
    "Le groupe CHU souhaite exploiter ses données de soins, de satisfaction et de décès afin de "
    "piloter l'activité médicale. La première consultation des utilisateurs (praticiens, chefs "
    "d'établissement) a permis d'identifier 8 besoins d'analyse, qui fixent les axes et les mesures "
    "de l'entrepôt :"
)
add_table(
    ["#", "Besoin d'analyse imposé"],
    [
        ["B1", "Taux de consultation des patients dans un établissement X sur une période Y"],
        ["B2", "Taux de consultation des patients par rapport à un diagnostic X sur une période Y"],
        ["B3", "Taux global d'hospitalisation des patients dans une période donnée Y"],
        ["B4", "Taux d'hospitalisation des patients par rapport à des diagnostics sur une période donnée"],
        ["B5", "Taux d'hospitalisation par sexe, par âge"],
        ["B6", "Taux de consultation par professionnel"],
        ["B7", "Nombre de décès par localisation (région) et sur l'année 2019"],
        ["B8", "Taux global de satisfaction par région sur l'année 2020"],
    ],
    widths=[0.5, 6.5],
)

doc.add_heading("2.2 Périmètre du Livrable 1", level=2)
doc.add_paragraph("Le présent livrable est de nature CONCEPTUELLE. Il comprend :")
bullets([
    "le modèle de données (faits, dimensions, grains, mesures) — axes d'analyse ;",
    "la description des jobs d'alimentation (ETL) ;",
    "la description de l'architecture de l'entrepôt.",
])
doc.add_paragraph("Sont explicitement renvoyés aux livrables suivants :")
bullets([
    "Livrable 2 : implémentation physique (scripts de création/chargement), partitionnement et "
    "bucketing, vérification des données, mesure et optimisation des performances ;",
    "Livrable 3 : tableaux de bord, restitution et storytelling.",
])

doc.add_heading("2.3 Équipe et responsabilités", level=2)
add_table(
    ["Membre", "Fait principal", "Source", "Responsabilités"],
    [
        ["Julian", "Fait_Consultation", "PostgreSQL", "Modélisation, ETL, dimension Professionnel"],
        ["Chloé", "Fait_Hospitalisation", "PostgreSQL / CSV", "Modélisation, ETL, dimension Diagnostic"],
        ["Matthieu", "Fait_Satisfaction", "Fichiers plats", "Modélisation, ETL, agrégations région"],
        ["Maxime", "Fait_Deces", "INSEE", "Modélisation, ETL, volumétrie"],
        ["Commun", "Dimensions + architecture", "Tous", "Temps, Patient, Établissement, Géographie, sécurité"],
    ],
    widths=[1.0, 1.8, 1.5, 2.7],
)

# =========================================================================
# 3. ARCHITECTURE MÉDAILLON
# =========================================================================
doc.add_heading("3. Architecture de l'entrepôt (médaillon Bronze / Silver / Gold)", level=1)
doc.add_paragraph(
    "L'entrepôt suit une architecture en médaillon à trois zones, adaptée à une chaîne batch sur "
    "données historiques (dumps statiques)."
)
add_table(
    ["Zone", "Rôle", "Contenu"],
    [
        ["Bronze (ingestion)", "Copie brute des sources sur HDFS",
         "Aucune transformation ; permet reprise et audit ; conservation historique"],
        ["Silver (nettoyage + anonymisation)", "Données dédupliquées, formatées, pseudonymisées",
         "Pseudonymisation patient (SHA-256), suppression colonnes sensibles, validation qualité"],
        ["Gold (entrepôt)", "Modèle en constellation (faits + dimensions)",
         "Partitionnement par année, bucketing sur clé de jointure, format Parquet (Snappy)"],
    ],
    widths=[2.0, 2.2, 2.8],
)
doc.add_paragraph("Flux global :", style="Intense Quote")
doc.add_paragraph(
    "Sources  →  BRONZE (brut HDFS)  →  SILVER (nettoyé + pseudonymisé)  →  GOLD (schéma en "
    "constellation Hive)  →  Restitution BI (Livrable 3)."
)
doc.add_paragraph(
    "Exemple (hospitalisations) : Hospitalisations.csv → Bronze (table externe) → Silver "
    "(déduplication, parsing des dates, validation) → pseudonymisation de l'identifiant patient → "
    "chargement de Fait_Hospitalisation (partitionné par année, bucketé par établissement)."
)

# =========================================================================
# 4. SOURCES ET VOLUMÉTRIE
# =========================================================================
doc.add_heading("4. Sources de données et volumétrie", level=1)
doc.add_heading("4.1 Vue d'ensemble", level=2)
add_table(
    ["Source", "Type", "Taille", "Volumétrie", "Sujet"],
    [
        ["PostgreSQL", "SGBD", "~45 Mo", "Consultations + hospitalisations", "Soins médico-administratifs"],
        ["CSV Établissements", "Fichiers plats", "~203 Mo", "Référentiel national", "Établissements / professionnels"],
        ["Fichiers Satisfaction", "Fichiers plats", "~16 Mo", "Campagnes eSATIS par année", "Notes de satisfaction"],
        ["INSEE Décès", "Fichier plat", "~1,9 Go", "≈ 25 millions de lignes", "Répertoire national des décès"],
    ],
    widths=[1.6, 1.3, 1.0, 1.9, 1.7],
)
doc.add_paragraph(
    "Justification de l'architecture distribuée : la source INSEE des décès (≈ 1,9 Go, ≈ 25 millions "
    "de lignes) à elle seule justifie le recours à HDFS + Hive ; l'accumulation historique "
    "(plusieurs années) et le volume total (≈ 2,1 Go) appellent une solution scalable en mode batch."
)

doc.add_heading("4.2 Détail des sources", level=2)

doc.add_heading("4.2.1 PostgreSQL — soins médico-administratifs", level=3)
doc.add_paragraph(
    "Base relationnelle (~45 Mo) contenant les consultations (colonnes : num_consultation, id_mut, "
    "id_patient, id_prof_sante, code_diag, motif, date_consultation, heure_debut, heure_fin) et les "
    "hospitalisations. Les identifiants patients y sont pseudonymisés dès la zone Silver."
)
doc.add_heading("4.2.2 CSV — établissements de santé", level=3)
doc.add_paragraph(
    "Référentiel des établissements français (dossier ≈ 203 Mo : établissements, professionnels, "
    "activité). Champs exploités : identifiant FINESS géographique, libellé, région, département. "
    "Sert d'alimentation aux dimensions Établissement et, indirectement, Géographie."
)
doc.add_heading("4.2.3 Fichiers plats — satisfaction (eSATIS)", level=3)
doc.add_paragraph(
    "Campagnes annuelles eSATIS (≈ 16 Mo). Chaque ligne porte un FINESS géographique, une région et "
    "un score global ajusté ; la source est déjà agrégée par établissement (aucun patient individuel). "
    "Table de fait : Fait_Satisfaction (campagne de référence : 2020 pour B8)."
)
doc.add_heading("4.2.4 INSEE — répertoire des décès", level=3)
doc.add_paragraph(
    "Fichier INSEE le plus volumineux (≈ 1,9 Go, ≈ 25 millions de lignes). Données d'état civil : "
    "nom, prénom, sexe, date et lieu de naissance, date de décès, code commune du lieu de décès, "
    "numéro d'acte. Cette source ne contient AUCUN identifiant du système de soins ni cause "
    "médicale : elle alimente Fait_Deces de manière autonome (axe région / sexe / âge), sans "
    "rapprochement avec les patients hospitalisés."
)

# =========================================================================
# 5. DIMENSIONS CONFORMES
# =========================================================================
doc.add_heading("5. Modèle conceptuel et dimensions conformes", level=1)
doc.add_paragraph(
    "Le modèle décisionnel adopté est un schéma en CONSTELLATION : 4 tables de faits (en vert) "
    "partagent 6 dimensions conformes (en bleu). La figure 1 en donne la vue d'ensemble ; les "
    "attributs détaillés des dimensions sont décrits ci-après (§5.1 à §5.6) et ceux des faits en "
    "section 6."
)
_fig = doc.add_paragraph()
_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fig.add_run().add_picture(
    "/Users/juliankerignard/Documents/CESI/BlocBigData/Projet/docs/mcd_constellation.png",
    width=Inches(6.6),
)
_cap = doc.add_paragraph()
_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
_cr = _cap.add_run("Figure 1 — Schéma en constellation : 4 faits, 6 dimensions conformes partagées (MCD)")
_cr.italic = True
_cr.font.size = Pt(9)
_cr.font.color.rgb = GREY
doc.add_paragraph()
doc.add_paragraph(
    "Les 6 dimensions suivantes sont modélisées une seule fois et partagées par les tables de faits, "
    "afin de garantir la cohérence des axes d'analyse (dimensions conformes)."
)

doc.add_heading("5.1 Dim_Temps", level=2)
doc.add_paragraph("Axe temporel commun aux 4 faits. Clé technique entière au format AAAAMMJJ.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["date_id", "INT", "Clé technique AAAAMMJJ (ex. 20190101)"],
        ["jour", "INT", "Jour du mois"],
        ["mois", "INT", "Mois (1-12)"],
        ["libelle_mois", "STRING", "Libellé du mois (Janvier … Décembre)"],
        ["trimestre", "INT", "Trimestre (1-4)"],
        ["annee", "INT", "Année (plage générée 2015-2023)"],
        ["jour_semaine", "STRING", "Libellé du jour (Lundi … Dimanche)"],
    ],
    widths=[1.6, 1.0, 4.4],
)

doc.add_heading("5.2 Dim_Patient", level=2)
doc.add_paragraph("Pseudonymisée (SHA-256). Aucun identifiant patient en clair n'est conservé.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["patient_id", "STRING", "Hash SHA-256 (64 caractères hexadécimaux) — pseudonymisé"],
        ["sexe", "STRING", "M ou F"],
        ["tranche_age", "STRING", "0-19, 20-39, 40-59, 60-74, 75-84, 85+"],
        ["region_residence", "STRING", "Région de résidence (code INSEE)"],
    ],
    widths=[1.8, 1.0, 4.2],
)
doc.add_paragraph(
    "Note : alimentée depuis le Bronze PostgreSQL (consultations / hospitalisations). Le hash est "
    "produit par le job de pseudonymisation (section 9.2) afin d'être identique d'un fait à l'autre.",
    style="Intense Quote",
)

doc.add_heading("5.3 Dim_Diagnostic", level=2)
doc.add_paragraph("Référentiel des codes CIM-10. La généralisation par chapitre soutient B2 et B4.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["diag_id", "STRING", "Clé technique (code CIM-10 observé)"],
        ["code_cim10", "STRING", "Code diagnostic CIM-10"],
        ["libelle", "STRING", "Libellé (NULL tant qu'aucun référentiel de libellés n'est fourni)"],
        ["chapitre_cim10", "STRING", "Catégorie généralisée (1re lettre du code) — minimisation RGPD"],
    ],
    widths=[1.6, 1.0, 4.4],
)

doc.add_heading("5.4 Dim_Etablissement", level=2)
doc.add_paragraph("Référentiel des établissements (clé FINESS géographique).")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["etab_id", "STRING", "FINESS géographique (site) — clé"],
        ["nom_etab", "STRING", "Nom de l'établissement"],
        ["type_etab", "STRING", "Type (MCO / SSR / HAD / PSY)"],
        ["region", "STRING", "Région"],
        ["departement", "STRING", "Département"],
    ],
    widths=[1.6, 1.0, 4.4],
)

doc.add_heading("5.5 Dim_Professionnel", level=2)
doc.add_paragraph(
    "Dimension requise par le besoin B6 (taux de consultation par professionnel). Alimentée depuis "
    "le Bronze PostgreSQL (professionnels + spécialités)."
)
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["prof_id", "STRING", "Identifiant du professionnel (ADELI)"],
        ["specialite", "STRING", "Spécialité"],
        ["categorie_prof", "STRING", "Catégorie professionnelle"],
        ["code_specialite", "STRING", "Code de spécialité"],
    ],
    widths=[1.7, 1.0, 4.3],
)

doc.add_heading("5.6 Dim_Geographie", level=2)
doc.add_paragraph(
    "Axe régional commun aux décès (B7) et à la satisfaction (B8). Garantit que les deux faits "
    "pointent vers les mêmes régions (découpage 2016)."
)
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["geo_id", "STRING", "Clé technique (code région)"],
        ["code_region", "STRING", "Code région INSEE"],
        ["region", "STRING", "Nom de la région"],
        ["code_departement", "STRING", "Code département (grain optionnel)"],
        ["departement", "STRING", "Nom du département (grain optionnel)"],
    ],
    widths=[1.7, 1.0, 4.3],
)

# =========================================================================
# 6. TABLES DE FAITS
# =========================================================================
doc.add_heading("6. Tables de faits", level=1)
doc.add_paragraph(
    "Quatre tables de faits, partitionnées par année et stockées en Parquet. Les clés de jointure "
    "(date_id, patient_id, prof_id, diag_id, etab_id, geo_id) sont conformes aux dimensions."
)

doc.add_heading("6.1 Fait_Consultation (Julian)", level=2)
doc.add_paragraph("Source : PostgreSQL. Grain : une ligne = une consultation. Besoins : B2, B6 (et B1, cf. note).")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["consultation_key", "BIGINT", "Clé technique (surrogate, remplace num_consultation)"],
        ["date_id", "INT", "FK → Dim_Temps (date de consultation)"],
        ["patient_id", "STRING", "FK → Dim_Patient (pseudonymisé)"],
        ["prof_id", "STRING", "FK → Dim_Professionnel"],
        ["diag_id", "STRING", "FK → Dim_Diagnostic"],
        ["nb_consultation", "INT", "Mesure (grain = 1)"],
        ["duree_minutes", "DOUBLE", "Mesure : durée calculée (heure_fin − heure_debut)"],
    ],
    widths=[1.9, 1.0, 4.1],
)
doc.add_paragraph("Mesures : nb_consultation, duree_minutes.  Dimensions : Temps, Patient, Professionnel, Diagnostic.")
doc.add_paragraph(
    "Note B1 : la source des consultations ne porte pas d'identifiant d'établissement (id_mut "
    "désigne la mutuelle, non un FINESS). Le besoin B1 « par établissement » est donc traité au "
    "niveau de l'établissement unique du périmètre ; l'axe Établissement reste disponible pour les "
    "hospitalisations.",
    style="Intense Quote",
)

doc.add_heading("6.2 Fait_Hospitalisation (Chloé)", level=2)
doc.add_paragraph("Source : PostgreSQL / CSV. Grain : une ligne = un séjour. Besoins : B3, B4, B5.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["hosp_key", "BIGINT", "Clé technique (surrogate, remplace num_hospitalisation)"],
        ["date_id", "INT", "FK → Dim_Temps (date d'entrée)"],
        ["patient_id", "STRING", "FK → Dim_Patient (pseudonymisé) → axe sexe/âge (B5)"],
        ["etab_id", "STRING", "FK → Dim_Etablissement"],
        ["diag_id", "STRING", "FK → Dim_Diagnostic"],
        ["nb_hospitalisation", "INT", "Mesure (grain = 1)"],
        ["duree_sejour", "INT", "Mesure : durée du séjour en jours (fournie par la source)"],
    ],
    widths=[2.0, 1.0, 4.0],
)
doc.add_paragraph("Mesures : nb_hospitalisation, duree_sejour.  Dimensions : Temps, Patient, Établissement, Diagnostic.")
doc.add_paragraph(
    "Note : la durée de séjour est fournie directement par la source (champ jour_hospitalisation) ; "
    "elle n'est pas recalculée par soustraction de dates. La source ne fournit ni type de séjour, "
    "ni information de décès en cours de séjour : ces axes ne sont pas modélisés.",
    style="Intense Quote",
)

doc.add_heading("6.3 Fait_Satisfaction (Matthieu)", level=2)
doc.add_paragraph("Source : fichiers plats eSATIS. Grain : une ligne = un établissement × campagne. Besoin : B8.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["satisfaction_key", "BIGINT", "Clé technique (surrogate)"],
        ["date_id", "INT", "FK → Dim_Temps (campagne, grain annuel AAAA0101)"],
        ["etab_id", "STRING", "FK → Dim_Etablissement"],
        ["geo_id", "STRING", "FK → Dim_Geographie (région) — axe commun avec B7"],
        ["note_satisfaction", "DECIMAL(3,1)", "Mesure : score global ajusté (échelle /10)"],
    ],
    widths=[2.0, 1.2, 3.8],
)
doc.add_paragraph("Mesure : note_satisfaction (non additive, agrégée par moyenne).  Dimensions : Temps, Établissement, Géographie.")

doc.add_heading("6.4 Fait_Deces (Maxime)", level=2)
doc.add_paragraph("Source : INSEE (≈ 25 M lignes). Grain : une ligne = un décès. Besoin : B7.")
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["deces_key", "BIGINT", "Clé technique (surrogate)"],
        ["date_id", "INT", "FK → Dim_Temps (date de décès)"],
        ["geo_id", "STRING", "FK → Dim_Geographie (région du lieu de décès)"],
        ["sexe", "STRING", "Dimension dégénérée (M / F)"],
        ["tranche_age", "STRING", "Dimension dégénérée (âge au décès, par tranche)"],
        ["nb_deces", "INT", "Mesure (grain = 1)"],
    ],
    widths=[1.7, 1.0, 4.3],
)
doc.add_paragraph("Mesure : nb_deces.  Dimensions : Temps, Géographie + sexe/tranche_age (dégénérées).")
doc.add_paragraph(
    "Note : la source INSEE ne comportant ni identifiant patient du système de soins, ni "
    "établissement, ni cause médicale, le fait n'est rattaché ni à Dim_Patient ni à "
    "Dim_Etablissement. Le sexe et la tranche d'âge sont portés en dimensions dégénérées. La région "
    "(geo_id) est dérivée du code commune du lieu de décès via le référentiel département → région.",
    style="Intense Quote",
)

doc.add_heading("6.5 Matrice de couverture des besoins", level=2)
doc.add_paragraph("Chaque besoin imposé est rattaché à un fait et à ses dimensions d'analyse.")
add_table(
    ["Besoin", "Fait", "Dimensions d'analyse", "Statut"],
    [
        ["B1 consult / établissement / période", "Fait_Consultation", "Temps", "Limité (cf. note 6.1)"],
        ["B2 consult / diagnostic / période", "Fait_Consultation", "Diagnostic, Temps", "Couvert"],
        ["B3 hospi globale / période", "Fait_Hospitalisation", "Temps", "Couvert"],
        ["B4 hospi / diagnostic / période", "Fait_Hospitalisation", "Diagnostic, Temps", "Couvert"],
        ["B5 hospi / sexe / âge", "Fait_Hospitalisation", "Patient (sexe, tranche_age)", "Couvert"],
        ["B6 consult / professionnel", "Fait_Consultation", "Professionnel", "Couvert"],
        ["B7 décès / région / 2019", "Fait_Deces", "Géographie, Temps", "Couvert"],
        ["B8 satisfaction / région / 2020", "Fait_Satisfaction", "Géographie, Établissement, Temps", "Couvert"],
    ],
    widths=[2.4, 1.7, 2.1, 1.1],
)

# =========================================================================
# 7. ARCHITECTURE ET OUTILS
# =========================================================================
doc.add_heading("7. Architecture et outils technologiques", level=1)
doc.add_heading("7.1 Stack retenu", level=2)
add_table(
    ["Couche", "Outil", "Justification"],
    [
        ["Stockage", "HDFS (local)", "Distribué, tolérant aux pannes, scalable, open-source"],
        ["Ingestion", "hdfs dfs -put + COPY", "Suffisant pour ≈ 2 Go, simplifie la stack"],
        ["ETL / traitement", "HiveQL (batch)", "Partitionnement et bucketing natifs, SQL standard"],
        ["Entrepôt", "Hive (local)", "Métastore, connecteurs BI (JDBC/ODBC)"],
        ["Format", "Parquet (Snappy)", "Colonnaire, compressé, lectures analytiques rapides"],
        ["Sécurité", "SHA-256 + sel", "Pseudonymisation des patients (conformité RGPD)"],
        ["Restitution", "Power BI / Tableau", "Connexion à Hive, tableaux de bord (Livrable 3)"],
    ],
    widths=[1.5, 1.8, 3.7],
)
doc.add_heading("7.2 Justification des choix", level=2)
doc.add_paragraph("Spark écarté : à ≈ 2 Go, un moteur distribué est surdimensionné ; HiveQL réalise l'ETL directement en SQL.")
doc.add_paragraph("Mode batch (pas de streaming) : les sources sont des dumps statiques historiques ; le bucketing impose d'ailleurs une écriture par lots.")
doc.add_paragraph("Hive retenu : partitionnement/bucketing pour la performance, SQL connu du métier, métastore pour la gouvernance, connecteurs BI.")

# =========================================================================
# 8. JOBS D'ALIMENTATION ETL
# =========================================================================
doc.add_heading("8. Jobs d'alimentation ETL", level=1)
doc.add_heading("8.1 ETL plutôt qu'ELT", level=2)
doc.add_paragraph(
    "L'approche ETL (transformer avant de charger) est retenue car la pseudonymisation des données "
    "de santé doit être appliquée en amont (RGPD, article 9). Un ELT exposerait des identifiants "
    "patients en zone de staging. L'ETL permet en outre de valider la qualité et de rejeter les "
    "lignes invalides avant insertion dans l'entrepôt."
)

doc.add_heading("8.2 Architecture ETL standardisée (5 étapes)", level=2)
add_table(
    ["Étape", "Traitement"],
    [
        ["1. Extraction", "Lecture de la source, validation de schéma/encodage, dépôt en zone Bronze (copie brute)"],
        ["2. Nettoyage", "Déduplication (clé métier), gestion des NULL, standardisation des formats, contrôles de plage → Silver"],
        ["3. Transformation + pseudonymisation", "Pseudonymisation de l'identifiant patient (SHA-256 + sel), suppression des colonnes sensibles, dérivations (durée, tranche d'âge, région) → Silver"],
        ["4. Chargement", "Insertion dans la table de fait, partitionnement par année, bucketing, compression Parquet/Snappy → Gold"],
        ["5. Qualité & audit", "Réconciliation des volumes (source vs cible), contrôle des clés orphelines, table de rejets"],
    ],
    widths=[1.9, 5.1],
)

doc.add_heading("8.3 Jobs spécifiques", level=2)

doc.add_heading("8.3.1 ETL Consultation (Julian)", level=3)
doc.add_paragraph("Source : PostgreSQL → Fait_Consultation. Dimensions : Temps, Patient, Professionnel, Diagnostic.")
bullets([
    "déduplication défensive sur num_consultation ; rejet des lignes sans clé obligatoire ;",
    "pseudonymisation de id_patient (SHA-256 + sel) ; suppression du motif (texte libre, risque PII) ;",
    "calcul de duree_minutes (heure_fin − heure_debut) ; génération du surrogate consultation_key ;",
    "code diagnostic manquant → 'UNKNOWN' ; filtre sur la plage temporelle 2015-2023.",
])

doc.add_heading("8.3.2 ETL Hospitalisation (Chloé)", level=3)
doc.add_paragraph("Source : CSV (séparateur ';') → Fait_Hospitalisation. Dimensions : Temps, Patient, Établissement, Diagnostic.")
bullets([
    "déduplication (id_patient, num_hospitalisation, date_entree) ; rejet des clés manquantes ;",
    "parsing de la date d'entrée (JJ/MM/AAAA) ; durée de séjour reprise directement de jour_hospitalisation ;",
    "pseudonymisation de id_patient (même formule que les consultations → cohérence inter-faits) ;",
    "suppression de suite_diagnostic_consultation (texte libre) ; surrogate hosp_key.",
])
doc.add_paragraph(
    "Aucun rapprochement avec le fichier des décès n'est effectué (absence de clé commune).",
    style="Intense Quote",
)

doc.add_heading("8.3.3 ETL Satisfaction (Matthieu)", level=3)
doc.add_paragraph("Source : fichiers plats eSATIS → Fait_Satisfaction. Dimensions : Temps, Établissement, Géographie.")
bullets([
    "transcodage de l'encodage source vers UTF-8 ; normalisation du score (virgule → point, échelle /10) ;",
    "date_id dérivée de l'année de campagne (AAAA0101) ; résolution de geo_id via le libellé de région ;",
    "rejet des scores non diffusés (NULL) ou hors plage ; table de rejets pour l'audit qualité.",
])

doc.add_heading("8.3.4 ETL Décès (Maxime)", level=3)
doc.add_paragraph("Source : INSEE (≈ 25 M lignes) → Fait_Deces. Dimensions : Temps, Géographie.")
bullets([
    "suppression de nom, prénom et numéro d'acte (PII directe) ;",
    "codage du sexe 1/2 → M/F ; calcul de la tranche d'âge (année de décès − année de naissance) ;",
    "dérivation de la région (geo_id) depuis le code commune du lieu de décès (DOM sur 3 caractères, "
    "Corse 2A/2B) via le référentiel département → région ;",
    "filtre qualité sur la date de décès ; focus sur l'année 2019 pour B7.",
])

# =========================================================================
# 9. SÉCURITÉ ET ANONYMISATION
# =========================================================================
doc.add_heading("9. Sécurité et anonymisation", level=1)
doc.add_heading("9.1 Cadre", level=2)
doc.add_paragraph(
    "Les données de santé sont des données sensibles (RGPD, article 9 ; secret médical). L'approche "
    "retenue : pseudonymisation obligatoire des patients, minimisation (suppression des champs "
    "inutiles), généralisation (chapitres CIM-10, tranches d'âge), et restitution sur agrégats."
)
doc.add_heading("9.2 Pseudonymisation", level=2)
doc.add_paragraph(
    "L'identifiant patient est remplacé, dès la zone Silver, par un pseudonyme calculé avec un "
    "algorithme SHA-256 et un sel propre à chaque patient :"
)
bullets([
    "sel_patient = SHA-256(id_patient || graine_secrète) — déterministe et indépendant de l'année ;",
    "patient_id = SHA-256(id_patient || clé_maître || sel_patient).",
])
doc.add_paragraph(
    "Le sel étant dérivé du seul identifiant patient (et non de l'année de traitement), un même "
    "patient obtient toujours le même pseudonyme, quelle que soit l'année : le suivi inter-années "
    "reste possible. La correspondance identifiant ↔ pseudonyme est conservée dans une table de "
    "mapping séparée, à accès restreint."
)
doc.add_paragraph(
    "Pseudonymisation vs anonymisation : la pseudonymisation est réversible via la table de mapping "
    "sécurisée (elle convient à un entrepôt décisionnel) ; l'anonymisation, elle, rend la "
    "ré-identification impossible et s'applique aux restitutions publiques agrégées.",
    style="Intense Quote",
)
doc.add_heading("9.3 Règles par source", level=2)
add_table(
    ["Source", "Règles appliquées"],
    [
        ["Hospitalisations", "Num_Hospitalisation → surrogate ; Id_patient → pseudonymisé ; identifiant_organisation → conservé (etab_id) ; Code_diagnostic → généralisé (chapitre) ; Suite_diagnostic → supprimé"],
        ["Consultations", "Id_patient → pseudonymisé ; Motif → supprimé (texte libre) ; code diagnostic → généralisé via Dim_Diagnostic"],
        ["Décès (INSEE)", "Nom / prénom / numéro d'acte → supprimés ; sexe → M/F ; lieu de décès → généralisé en région"],
        ["Satisfaction", "Source déjà agrégée par établissement → aucune donnée patient à pseudonymiser"],
    ],
    widths=[1.6, 5.4],
)
doc.add_paragraph(
    "Grain temporel : les faits conservent le grain JOUR (date de l'événement), choix adapté aux "
    "analyses par période (B1-B8). La satisfaction et les décès sont par nature agrégés à l'année / "
    "à la campagne.",
    style="Intense Quote",
)
doc.add_heading("9.4 Contrôle d'accès (principe)", level=2)
doc.add_paragraph(
    "Le modèle d'accès distingue les profils administrateur (accès au mapping), analyste/BI (données "
    "pseudonymisées, lecture seule) et praticien/gestionnaire (tableaux de bord agrégés). La mise en "
    "œuvre technique des rôles relève des Livrables 2/3."
)

# =========================================================================
# 10. EXIGENCES NON FONCTIONNELLES (CIBLES)
# =========================================================================
doc.add_heading("10. Exigences non fonctionnelles (cibles de conception)", level=1)
doc.add_paragraph(
    "Les objectifs ci-dessous orientent la conception ; leur mesure et leur validation sont "
    "réalisées au Livrable 2 (performance) et au Livrable 3 (restitution)."
)
add_table(
    ["Axe", "Cible visée"],
    [
        ["Performance", "Partitionnement par année (élagage I/O) et bucketing sur la clé de jointure dominante"],
        ["Scalabilité", "Tenue de la volumétrie décès (≈ 25 M lignes) ; archivage par partitions annuelles"],
        ["Sécurité", "Pseudonymisation SHA-256 + sel ; minimisation et généralisation des données sensibles"],
        ["Qualité", "Réconciliation des volumes source/cible, contrôle des clés, table de rejets par job"],
        ["Maintenabilité", "Pipeline 100 % HiveQL (open-source), modèle documenté et versionné"],
    ],
    widths=[1.5, 5.5],
)

# =========================================================================
# 11. GLOSSAIRE
# =========================================================================
doc.add_heading("11. Glossaire", level=1)
add_table(
    ["Terme", "Définition"],
    [
        ["CIM-10", "Classification Internationale des Maladies (OMS)"],
        ["FINESS", "Fichier national des établissements sanitaires et sociaux"],
        ["RGPD", "Règlement Général sur la Protection des Données (UE 2016/679)"],
        ["SHA-256", "Fonction de hachage cryptographique (digest de 64 caractères hexadécimaux)"],
        ["Constellation", "Modèle décisionnel : plusieurs faits partageant des dimensions conformes"],
        ["ETL", "Extract – Transform – Load (pipeline d'alimentation)"],
        ["Dimension dégénérée", "Attribut d'analyse porté directement par le fait (ex. sexe, tranche d'âge)"],
        ["HDFS", "Hadoop Distributed File System"],
        ["Hive", "Moteur SQL batch sur HDFS"],
        ["Parquet", "Format de stockage colonnaire compressé"],
        ["Pseudonymisation", "Substitution réversible (via mapping sécurisé) d'un identifiant"],
        ["Anonymisation", "Transformation rendant la ré-identification impossible"],
    ],
    widths=[1.8, 5.2],
)

doc.save(OUT)
print("Document généré :", OUT)
print("Sections :", sum(1 for p in doc.paragraphs if p.style.name == "Heading 1"))
print("Tableaux :", len(doc.tables))
